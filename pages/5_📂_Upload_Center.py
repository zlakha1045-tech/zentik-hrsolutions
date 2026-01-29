import streamlit as st
from supabase import create_client
from styles import load_css, hero_section
import requests
import time
import mimetypes
import pdfplumber
import docx

st.set_page_config(page_title="Bulk Candidate Upload", layout="wide")
load_css()

# --- CONNECT DB ---
try:
    supabase = create_client(st.secrets["supabase"]["url"], st.secrets["supabase"]["key"])
except:
    st.error("Missing Secrets.")
    st.stop()

N8N_WEBHOOK = st.secrets["n8n"]["webhook_url"]

# --- IMPROVED TEXT EXTRACTION ---
def read_pdf(file):
    text = ""
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                extract = page.extract_text()
                if extract: text += extract + "\n"
    except Exception as e:
        return f"Error reading PDF: {e}"
    return text

def read_docx(file):
    """
    Extracts text from Paragraphs, Tables, Headers, AND Footers.
    """
    try:
        doc = docx.Document(file)
        full_text = []
        
        # 1. Extract Main Body Paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # 2. Extract Tables (Layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
        
        # 3. Extract Headers & Footers (Where contact info hides!)
        for section in doc.sections:
            # Check Header
            for header_para in section.header.paragraphs:
                full_text.append(header_para.text)
            # Check Footer
            for footer_para in section.footer.paragraphs:
                full_text.append(footer_para.text)
                        
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading Word Doc: {e}"

# --- MAIN APP UI ---
hero_section("assets/login_hero.jpg", "Candidate Portal", "Fast-Track Bulk Upload.")

st.subheader("📂 Resume Upload")

# 1. SELECT JOB
try:
    jobs = supabase.table("jobs").select("id, title, department").eq("status", "Active").execute().data
except:
    st.stop()

if not jobs:
    st.info("⚠️ No Active Jobs found.")
    st.stop()

job_map = {f"{j['title']} ({j['department']})": j for j in jobs}
selected_label = st.selectbox("Assign to Role:", list(job_map.keys()))
selected_job = job_map[selected_label]

st.divider()

# 2. BULK UPLOAD
uploaded_files = st.file_uploader("Drag resumes here (PDF or Docx)", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    if st.button(f"🚀 Fast-Process {len(uploaded_files)} Resumes", type="primary"):
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # --- A. UPLOAD LOOP (Fast) ---
        for i, uploaded_file in enumerate(uploaded_files):
            status_text.markdown(f"**Uploading ({i+1}/{len(uploaded_files)}):** `{uploaded_file.name}`")
            
            try:
                # 1. EXTRACT TEXT (Now handles Tables in Docx)
                if uploaded_file.name.lower().endswith(".pdf"):
                    resume_text = read_pdf(uploaded_file)
                else:
                    resume_text = read_docx(uploaded_file)
                
                # 2. UPLOAD FILE TO STORAGE
                file_ext = uploaded_file.name.split('.')[-1]
                timestamp = int(time.time())
                file_path = f"candidates/{selected_job['id']}/{timestamp}_{i}.{file_ext}"
                mime = mimetypes.guess_type(uploaded_file.name)[0]

                supabase.storage.from_("resumes").upload(file_path, uploaded_file.getvalue(), {"content-type": mime})
                resume_url = supabase.storage.from_("resumes").get_public_url(file_path)

                # 3. SAVE TO DB
                final_email = f"processing_{timestamp}_{i}@zentiklabs.com"
                cand_data = {
                    "job_id": selected_job['id'],
                    "name": uploaded_file.name, 
                    "email": final_email,
                    "resume_url": resume_url,
                    "status": "AI Analysis in Progress"
                }
                
                response = supabase.table("candidates").insert(cand_data).execute()
                new_row_id = response.data[0]['id']

                # 4. TRIGGER N8N (Fire and Forget)
                payload = {
                    "candidate_id": new_row_id,
                    "resume_text": resume_text, # Contains table text now
                    "resume_url": resume_url,
                    "job_id": selected_job['id']
                }
                try:
                    requests.post(N8N_WEBHOOK, json=payload, timeout=2) 
                except requests.exceptions.ReadTimeout:
                    pass 
                except Exception as e:
                    print(f"Webhook error: {e}")

            except Exception as e:
                st.error(f"❌ Error on {uploaded_file.name}: {e}")
            
            # Update Upload Progress
            progress_bar.progress((i + 1) / len(uploaded_files))
            
        # --- B. THE 20s BUFFER TIMER ---
        # This keeps the user on the page while n8n finishes in the background
        status_text.markdown("✅ **Uploads Complete! Finalizing AI Analysis...**")
        
        # We reuse the progress bar for the countdown
        buffer_seconds = 20
        for tick in range(buffer_seconds):
            time.sleep(1)
            remaining = buffer_seconds - tick
            status_text.markdown(f"⏳ **Syncing with AI Agent...** ({remaining}s remaining)")
        
        # --- C. FINISH ---
        status_text.markdown("## ✅ All Done!")
        st.balloons()
        st.success(f"Successfully processed {len(uploaded_files)} resumes.")
        
        st.info("👉 The AI has finished processing. Click **'Candidate Analysis'** in the sidebar to view the results.")